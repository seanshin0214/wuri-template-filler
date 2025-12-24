# -*- coding: utf-8 -*-
"""
WURI Template Auto-Filler
Data.docx의 데이터를 template.docx에 자동으로 채워넣는 스크립트

사용법:
    python fill_wuri_template.py

필요 라이브러리:
    pip install python-docx docxtpl
"""

from docx import Document
from docxtpl import DocxTemplate
import re
import os
import io
from datetime import datetime
import shutil

# ============================================================
# 1. Data.docx에서 데이터 추출
# ============================================================

def extract_data_from_docx(data_path):
    """Data.docx에서 구조화된 데이터를 추출합니다."""

    # 한글 경로 문제 해결: 바이너리로 읽어서 BytesIO로 전달
    with open(data_path, 'rb') as f:
        doc = Document(io.BytesIO(f.read()))
    full_text = "\n".join([para.text for para in doc.paragraphs])

    # 테이블에서도 데이터 추출
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join([cell.text for cell in row.cells])
            full_text += "\n" + row_text

    data = {}

    # ---- Writer's Profile ----
    data['university_name'] = extract_value(full_text, r'University name\s+(.+?)(?:\n|Full name)') or "MIT"
    data['writer_full_name'] = ""  # To be filled manually
    data['writer_relationship'] = ""  # To be filled manually
    data['writer_title'] = ""  # To be filled manually
    data['writer_email'] = ""  # To be filled manually

    # ---- Program Profile ----
    data['program_name'] = extract_value(full_text, r'Program name\s+(.+?)(?:\n|Category)') or "Amogy (Startup)"
    data['category'] = extract_value(full_text, r'Category\s+(C\d+:.+?)(?:\n|Champion|Abstract)') or "C7: University-Based Entrepreneurial Projects"
    data['champion'] = ""  # To be filled manually
    data['school_college'] = extract_value(full_text, r'School/College/Headquarters\s+(.+?)(?:\n|Summary)') or "1-UH: University Headquarters"

    # ---- Summary ----
    data['abstract'] = extract_section(full_text, r'Abstract\n(.+?)(?=Details of Program|$)')

    # ---- Planning ----
    data['long_term_goals'] = extract_section(full_text, r'Long-term Goals:\s*(.+?)(?=Short-term Targets:|$)')
    data['short_term_targets'] = extract_section(full_text, r'Short-term Targets:\s*(.+?)(?=Rationale:|$)')
    data['rationale'] = extract_section(full_text, r'Rationale:\s*(.+?)(?=Subject \(Leader\)|$)')

    data['initiators'] = extract_section(full_text, r'Initiator\(s\):\s*(.+?)(?=Champion\(s\):|$)')
    data['champions'] = extract_section(full_text, r'Champion\(s\):\s*(.+?)(?=Major team member|$)')
    data['team_members'] = extract_section(full_text, r'Major team member\(s\):\s*(.+?)(?=Environment|$)')

    data['nature_society'] = extract_section(full_text, r'Nature/Society:\s*(.+?)(?=Industry/Market:|$)')
    data['industry_market'] = extract_section(full_text, r'Industry/Market:\s*(.+?)(?=Citizen/Government:|$)')
    data['citizen_government'] = extract_section(full_text, r'Citizen/Government:\s*(.+?)(?=Resources|$)')

    data['human_resources'] = extract_section(full_text, r'Human resources:\s*(.+?)(?=Financial resources:|$)')
    data['financial_resources'] = extract_section(full_text, r'Financial resources:\s*(.+?)(?=Technological resources:|$)')
    data['technological_resources'] = extract_section(full_text, r'Technological resources:\s*(.+?)(?=Mechanism|$)')

    data['strategy'] = extract_section(full_text, r'Strategy:\s*(.+?)(?=Organization:|$)')
    data['organization'] = extract_section(full_text, r'Organization:\s*(.+?)(?=Culture:|$)')
    data['culture'] = extract_section(full_text, r'Culture:\s*(.+?)(?=2\. Doing|$)')

    # ---- Doing ----
    data['launch_date'] = extract_value(full_text, r'Launch date:\s*(.+?)(?:\n|Responsible)')
    data['responsible_org'] = extract_value(full_text, r'Responsible organization:\s*(.+?)(?:\n|Program content)')
    data['program_content'] = extract_section(full_text, r'Program content and process\n(.+?)(?=Key highlights|$)')
    data['content_highlights'] = extract_section(full_text, r'Content highlights:\s*(.+?)(?=Process highlights:|$)')
    data['process_highlights'] = extract_section(full_text, r'Process highlights:\s*(.+?)(?=Differences from|$)')
    data['differences'] = extract_section(full_text, r'Differences from traditional approaches:\s*(.+?)(?=Progress as of today:|$)')
    data['progress'] = extract_section(full_text, r'Progress as of today:\s*(.+?)(?=Problems in implementation:|$)')
    data['problems'] = extract_section(full_text, r'Problems in implementation:\s*(.+?)(?=Approaches to solve|$)')
    data['solutions'] = extract_section(full_text, r'Approaches to solve the problems:\s*(.+?)(?=Completion date|$)')
    data['completion_date'] = extract_value(full_text, r'Completion date:\s*(.+?)(?:\n|3\. Seeing)')

    # ---- Seeing ----
    data['impacts_students'] = extract_section(full_text, r'Impacts on students:\s*(.+?)(?=Impacts on professors|$)')
    data['impacts_professors'] = extract_section(full_text, r'Impacts on professors:\s*(.+?)(?=Impacts on university|$)')
    data['impacts_admin'] = extract_section(full_text, r'Impacts on university administration:\s*(.+?)(?=Responses from industry|$)')
    data['responses_industry'] = extract_section(full_text, r'Responses from industry/market:\s*(.+?)(?=Responses from citizen|$)')
    data['responses_citizen'] = extract_section(full_text, r'Responses from citizen/government:\s*(.+?)(?=Measurable output|$)')
    data['measurable_output'] = extract_section(full_text, r'Measurable output \(revenues\):\s*(.+?)(?=Measurable input|$)')
    data['measurable_input'] = extract_section(full_text, r'Measurable input \(expenses\):\s*(.+?)(?=Cost-benefit|$)')
    data['cost_benefit'] = extract_section(full_text, r'Cost-benefit analysis:\s*(.+?)(?=4\. Future Planning|$)')

    # ---- Future Planning ----
    data['future_planning'] = extract_section(full_text, r'Where does the project go from here\?:\s*(.+?)(?=5\. Addendum|$)')

    # ---- Addendum ----
    data['website_links'] = extract_section(full_text, r'Website links\n(.+?)(?=Supporting documents|$)')

    # Clean all values
    for key in data:
        if data[key]:
            data[key] = clean_text(data[key])
        else:
            data[key] = ""

    return data


def extract_value(text, pattern):
    """정규식으로 단일 값 추출"""
    match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
    if match:
        return match.group(1).strip()
    return None


def extract_section(text, pattern):
    """정규식으로 섹션 추출"""
    match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
    if match:
        return match.group(1).strip()
    return None


def clean_text(text):
    """텍스트 정리"""
    if not text:
        return ""
    # 여러 줄바꿈을 하나로
    text = re.sub(r'\n{3,}', '\n\n', text)
    # 앞뒤 공백 제거
    text = text.strip()
    return text


# ============================================================
# 2. 플레이스홀더가 있는 템플릿 생성 (최초 1회)
# ============================================================

def create_template_with_placeholders(original_template_path, output_path):
    """
    원본 template.docx를 읽고 플레이스홀더를 추가한 버전 생성

    주의: 이 함수는 템플릿 구조를 보존하면서 빈 셀에 플레이스홀더를 추가합니다.
    복잡한 테이블 구조에서는 수동 편집이 더 정확할 수 있습니다.
    """
    # 한글 경로 문제 해결
    with open(original_template_path, 'rb') as f:
        doc = Document(io.BytesIO(f.read()))

    # 플레이스홀더 매핑 (테이블 셀 텍스트 -> 플레이스홀더)
    placeholder_map = {
        'University name': '{{university_name}}',
        'Full name': '{{writer_full_name}}',
        'Relationship': '{{writer_relationship}}',
        'Official title': '{{writer_title}}',
        'Email address': '{{writer_email}}',
        'Phone number': '{{writer_phone}}',
        'Program name': '{{program_name}}',
        'Category': '{{category}}',
        'Champion': '{{champion}}',
        'Program Name': '{{program_name}}',
        'Abstract of Program': '{{abstract}}',
        'Long-term Goals': '{{long_term_goals}}',
        'Short-term Targets': '{{short_term_targets}}',
        'Rationale': '{{rationale}}',
        'Initiator(s)': '{{initiators}}',
        'Champion(s)': '{{champions}}',
        'Major team member(s)': '{{team_members}}',
        'Nature/Society': '{{nature_society}}',
        'Industry/Market': '{{industry_market}}',
        'Citizen/Government': '{{citizen_government}}',
        'Human resources': '{{human_resources}}',
        'Financial resources': '{{financial_resources}}',
        'Technological resources': '{{technological_resources}}',
        'Strategy (Weight/Sequence)': '{{strategy}}',
        'Organization': '{{organization}}',
        'Culture': '{{culture}}',
        'Launch date': '{{launch_date}}',
        'Responsible organization': '{{responsible_org}}',
        'Program content and process': '{{program_content}}',
        'Key highlights of the content/process': '{{content_highlights}}\n{{process_highlights}}',
        'Differences from traditional approaches': '{{differences}}',
        'Progress as of today': '{{progress}}',
        'Problems in implementation': '{{problems}}',
        'Approaches to solve the problems': '{{solutions}}',
        'Completion date, if completed': '{{completion_date}}',
        'Impacts on students': '{{impacts_students}}',
        'Impacts on professors': '{{impacts_professors}}',
        'Impacts on university administration': '{{impacts_admin}}',
        'Responses from industry/market': '{{responses_industry}}',
        'Responses from citizen/government': '{{responses_citizen}}',
        'Measurable output (revenues)': '{{measurable_output}}',
        'Measurable input (expenses)': '{{measurable_input}}',
        'Cost-benefit analysis for effectiveness': '{{cost_benefit}}',
        'Where does the project go from here?': '{{future_planning}}',
    }

    # 테이블의 각 셀 순회
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            for i, cell in enumerate(cells):
                cell_text = cell.text.strip()
                # 다음 셀이 비어있으면 플레이스홀더 추가
                if cell_text in placeholder_map and i + 1 < len(cells):
                    next_cell = cells[i + 1]
                    if not next_cell.text.strip() or next_cell.text.strip() == '':
                        next_cell.text = placeholder_map[cell_text]

    doc.save(output_path)
    print(f"플레이스홀더 템플릿 생성: {output_path}")
    return output_path


# ============================================================
# 3. 템플릿에 데이터 채우기
# ============================================================

def fill_template(template_path, data, output_path):
    """docxtpl을 사용해 템플릿에 데이터를 채웁니다."""

    # 한글 경로 문제 해결
    with open(template_path, 'rb') as f:
        doc = DocxTemplate(io.BytesIO(f.read()))
    doc.render(data)

    # 출력 파일 저장
    with open(output_path, 'wb') as f:
        doc.save(f)
    print(f"완성된 문서 저장: {output_path}")
    return output_path


# ============================================================
# 메인 실행
# ============================================================

def main():
    # 경로 설정 - 스크립트 파일과 같은 폴더에 모든 파일 배치
    script_dir = os.path.dirname(os.path.abspath(__file__))

    # 같은 폴더의 파일 사용 (복사본)
    data_path = os.path.join(script_dir, 'Data.docx')
    template_path = os.path.join(script_dir, 'template_original.docx')
    placeholder_template_path = os.path.join(script_dir, 'template_with_placeholders.docx')

    # 경로 확인
    print(f"Data 경로: {data_path}")
    print(f"Template 경로: {template_path}")
    print(f"파일 존재 확인: Data={os.path.exists(data_path)}, Template={os.path.exists(template_path)}")

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    output_path = os.path.join(script_dir, f'WURI_Filled_{timestamp}.docx')

    print("=" * 60)
    print("WURI Template Auto-Filler")
    print("=" * 60)

    # Step 1: Data.docx에서 데이터 추출
    print("\n[1/3] Data.docx에서 데이터 추출 중...")
    data = extract_data_from_docx(data_path)
    print(f"  - 추출된 필드 수: {len([k for k, v in data.items() if v])}")

    # 추출된 데이터 미리보기
    print("\n  [추출된 주요 데이터]")
    print(f"  - 대학명: {data.get('university_name', 'N/A')}")
    print(f"  - 프로그램명: {data.get('program_name', 'N/A')}")
    print(f"  - 카테고리: {data.get('category', 'N/A')}")

    # Step 2: 플레이스홀더 템플릿 생성 (없으면)
    print("\n[2/3] 플레이스홀더 템플릿 확인 중...")
    if not os.path.exists(placeholder_template_path):
        print("  - 플레이스홀더 템플릿 새로 생성...")
        create_template_with_placeholders(template_path, placeholder_template_path)
    else:
        print(f"  - 기존 템플릿 사용: {placeholder_template_path}")

    # Step 3: 템플릿에 데이터 채우기
    print("\n[3/3] 템플릿에 데이터 채우는 중...")
    fill_template(placeholder_template_path, data, output_path)

    print("\n" + "=" * 60)
    print("완료!")
    print(f"결과 파일: {output_path}")
    print("=" * 60)


if __name__ == "__main__":
    main()
